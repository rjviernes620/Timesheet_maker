from docx import Document
from re import findall
from datetime import datetime

class request_sheet:

    def __init__(self, sheet):
        self.sheet = sheet
        self.details = dict()

    def budget_code(self):
        for table in self.sheet.tables:
            for row in table.rows:
                if 'Budget Code' in row.cells[0].text:
                    return row.cells[1].text
        return None
    
    def shift_date(self):
        for table in self.sheet.tables:
            for row in table.rows:
                if row.cells[0].text == 'Shift Date(s) and Times':
                    long_date = row.cells[1].text
                    new = [int(num) for num in findall(r'\d+', long_date) if int(num) < 32 != datetime.now().year]
                    result = ''
                    for i in new:
                        result += str(i)
                    return result
        return None
    
    def start_time(self):
        for table in self.sheet.tables:
            for row in table.rows:
                if row.cells[0].text == 'Shift Date(s) and Times':
                    return "{:.2f}".format(float(row.cells[2].text.split('-')[0].replace(':', '.')))

        return None
    
    def break_time(self, start, end):
        shift_duration = float(end) - float(start)

        # Calculate total break time in hours (30 minutes for every 6 hours)
        total_break_time_hours = (shift_duration // 6) * 0.5

        if total_break_time_hours == 0:
            return 'NO BREAK'
        
        else:
            # Convert start time to total minutes
            start_hours, start_minutes = map(int, start.split('.'))
            start_total_minutes = start_hours * 60 + start_minutes

            # Calculate shift duration in minutes
            shift_duration_minutes = shift_duration * 60

            # Calculate total break time in minutes
            total_break_time_minutes = total_break_time_hours * 60

            # Calculate break start time in minutes to center the break
            break_start_minutes = start_total_minutes + (shift_duration_minutes - total_break_time_minutes) / 2

            # Ensure break_start_minutes and break_end_minutes stay within a 24-hour format (1440 minutes)
            break_start_minutes = break_start_minutes % 1440
            break_end_minutes = (break_start_minutes + total_break_time_minutes) % 1440

            # Convert break start and end times back to hours and minutes
            break_start_hours = int(break_start_minutes // 60)
            break_start_minutes = int(break_start_minutes % 60)
            break_end_hours = int(break_end_minutes // 60)
            break_end_minutes = int(break_end_minutes % 60)

            # Format break times
            break_start_formatted = f"{break_start_hours:02d}.{break_start_minutes:02d}"
            break_end_formatted = f"{break_end_hours:02d}.{break_end_minutes:02d}"

            return [float(break_start_formatted), float(break_end_formatted)]

    
    def end_time(self):
        for table in self.sheet.tables:
            for row in table.rows:
                if row.cells[0].text == 'Shift Date(s) and Times':
                    return "{:.2f}".format(float(row.cells[2].text.split('-')[1].replace(':', '.')))

        return None
    
    def shift_name(self):
        for table in self.sheet.tables:
            for row in table.rows:
                if 'Event / Activity' in row.cells[0].text:
                    return row.cells[1].text
        return None
    
    def pay_rate(self):
        for table in self.sheet.tables:
            for row in table.rows:
                if 'Rate of Pay' in row.cells[0].text:
                    return row.cells[1].text
        return None
    
    def collate_to_dict(self):
        self.details["shift_date"] = int(self.shift_date())
        self.details["start_time"] = self.start_time()
        self.details["start_break"] = self.break_time(self.start_time(), self.end_time())[0]
        self.details["end_break"] = self.break_time(self.start_time(), self.end_time())[1]
        self.details["end_time"] = self.end_time()
        self.details["pay_rate"] = self.pay_rate()
        self.details["budget_code"] = self.budget_code()
        self.details["shift_name"] = self.shift_name()
        print(self.details)
        return self.details
    
doc = Document('demo\\src\\main\\resources\\test1.docx')
request = request_sheet(doc)
request.collate_to_dict()
